"""Document generator service for creating .docx files."""
import json
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from bot.config import GENERATED_DOCS_DIR
from bot.locales import get_text
from bot.translator import translate_answers


def get_question_text(question: dict, language: str) -> str:
    """Get question text in specified language."""
    text = question.get('text', '')
    
    # Try to parse as JSON (new format with ru/en/ar)
    try:
        texts = json.loads(text)
        return texts.get(language, texts.get('en', text))
    except (json.JSONDecodeError, TypeError):
        # Old format - plain text
        return text


async def generate_questionnaire_docx(questionnaire_data: dict, questions: list, 
                                       language: str = "en") -> Path:
    """
    Generate a .docx document from questionnaire answers.
    
    NOTE: For Arabic (ar) language, answers are auto-translated to English
    and document is generated in English for operators.
    
    Args:
        questionnaire_data: Dictionary with questionnaire info and answers
        questions: List of question dictionaries
        language: Language code ('ru', 'en', or 'ar')
    
    Returns:
        Path to generated document
    """
    doc = Document()
    
    # NOTE: For Arabic questionnaires, generate document in English
    # with translated answers for operator convenience
    doc_language = 'en' if language == 'ar' else language
    is_arabic_source = language == 'ar'
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Title
    title_text = get_text(doc_language, "doc_title")
    if is_arabic_source:
        title_text += " (Translated from Arabic)"
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata section
    doc.add_paragraph()
    meta_para = doc.add_paragraph()
    
    respondent_name = questionnaire_data.get('respondent_name', 'N/A')
    respondent_username = questionnaire_data.get('respondent_username')
    invite_description = questionnaire_data.get('invite_description')
    completed_at = questionnaire_data.get('completed_at', datetime.now().isoformat())
    
    # Format date nicely
    try:
        dt = datetime.fromisoformat(completed_at)
        formatted_date = dt.strftime('%Y-%m-%d %H:%M')
    except:
        formatted_date = completed_at
    
    # Metadata labels based on document language
    if doc_language == "ru":
        respondent_label = "Респондент"
        date_label = "Дата"
        topic_label = "Тема"
    else:
        respondent_label = "Respondent"
        date_label = "Date"
        topic_label = "Topic"
    
    meta_text = f"{respondent_label}: {respondent_name}"
    if respondent_username:
        meta_text += f" (@{respondent_username})"
    meta_text += f"\n{date_label}: {formatted_date}"
    if invite_description:
        meta_text += f"\n{topic_label}: {invite_description}"
    if is_arabic_source:
        meta_text += "\n🌐 Original language: Arabic (auto-translated)"
    
    meta_para.add_run(meta_text)
    
    # Separator
    doc.add_paragraph("─" * 50)
    
    # Parse answers
    answers_json = questionnaire_data.get('answers_json', '{}')
    if isinstance(answers_json, str):
        answers = json.loads(answers_json)
    else:
        answers = answers_json
    
    # NOTE: Translate Arabic answers to English for operator
    if is_arabic_source and answers:
        answers = translate_answers(answers, source_lang='ar', target_lang='en')
    
    no_answer_text = get_text(doc_language, "no_answer")
    
    # Questions and Answers
    # NOTE: Questions shown in doc_language (English for Arabic questionnaires)
    for question in questions:
        q_key = question['key']
        q_text = get_question_text(question, doc_language)
        answer = answers.get(q_key, no_answer_text)
        
        # Question (bold)
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(q_text)
        q_run.bold = True
        q_run.font.size = Pt(11)
        
        # Answer
        a_para = doc.add_paragraph()
        a_run = a_para.add_run(answer)
        a_run.font.size = Pt(11)
        a_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        # Small spacing
        doc.add_paragraph()
    
    # Footer
    doc.add_paragraph("─" * 50)
    footer_para = doc.add_paragraph()
    footer_text = get_text(doc_language, "doc_generated_by")
    footer_run = footer_para.add_run(f"{footer_text} • {formatted_date}")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Generate filename
    safe_name = "".join(c for c in respondent_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    # NOTE: Arabic questionnaires saved with _ar_translated suffix
    lang_suffix = f"_{language}_translated" if is_arabic_source else f"_{language}"
    filename = f"questionnaire_{safe_name}{lang_suffix}_{timestamp}.docx"
    
    # Save document
    doc_path = GENERATED_DOCS_DIR / filename
    doc.save(doc_path)
    
    return doc_path
